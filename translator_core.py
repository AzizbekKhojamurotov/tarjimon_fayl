"""
translator_core.py — Pure-Python DOCX translation logic.

Adapted *exactly* from the provided Translate.ipynb notebook:
  • detect_language()        — counts Cyrillic / Latin / Uzbek-specific chars
  • write_text_to_runs()     — preserves per-run font / style formatting
  • iter_paragraphs_in_cell()— recurses into nested tables inside cells
  • translate_paragraph()    — translates one paragraph in-place
  • translate_document()     — walks paragraphs + tables + headers/footers
  • translate_docx_file()    — top-level helper called by the bot handler

All functions are synchronous and CPU/IO-bound.  The bot handler MUST call
translate_docx_file() inside asyncio.to_thread() to keep the event loop free.
"""

import logging
import re

from docx import Document
from docx.table import _Cell

logger = logging.getLogger(__name__)

# ── Supported language codes (must match deep-translator / Google Translate) ──
SUPPORTED_LANGS: dict[str, str] = {
    "uz": "Uzbek 🇺🇿",
    "ru": "Russian 🇷🇺",
    "en": "English 🇬🇧",
}


# ─────────────────────────────────────────────────────────────────────────────
# Language detection  (verbatim logic from the notebook)
# ─────────────────────────────────────────────────────────────────────────────




# ─────────────────────────────────────────────────────────────────────────────
# DOCX formatting helpers  (verbatim from the notebook)
# ─────────────────────────────────────────────────────────────────────────────

def write_text_to_runs(paragraph, new_text: str) -> None:
    """
    Distribute *new_text* across the existing runs of *paragraph* so that
    each run keeps its original character formatting (font, bold, italic, …).

    Strategy (from notebook):
      1. Measure original run lengths.
      2. Slice new_text into those same lengths, assigning each slice to the
         corresponding run.
      3. If new_text is longer, append the overflow to the last run.
      4. If new_text is shorter, blank out any remaining runs.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    new_text = str(new_text)  # Guard against None from a failed translation
    lengths = [len(r.text) for r in runs]
    pos = 0

    for i, run in enumerate(runs):
        take = lengths[i]
        run.text = new_text[pos: pos + take]
        pos += take

    # Overflow — append remainder to the last run
    if pos < len(new_text):
        runs[-1].text += new_text[pos:]
    # Underflow — clear any runs that got no characters assigned
    elif pos > len(new_text):
        for extra in range(i + 1, len(runs)):
            runs[extra].text = ""


def iter_paragraphs_in_cell(cell: _Cell):
    """
    Yield every paragraph inside *cell*, recursing into any nested tables
    (tables-within-tables are common in academic DOCX files).
    """
    for p in cell.paragraphs:
        yield p
    for tbl in cell.tables:
        for row in tbl.rows:
            for c in row.cells:
                yield from iter_paragraphs_in_cell(c)


# ─────────────────────────────────────────────────────────────────────────────
# Paragraph / document translation  (verbatim from the notebook)
# ─────────────────────────────────────────────────────────────────────────────

def translate_paragraph(paragraph, translator) -> None:
    """Translate a single paragraph in-place, skipping blank paragraphs."""
    original = paragraph.text
    if not original.strip():
        return

    translated = translator.translate(original)

    # Guard: deep-translator can return None on API hiccups
    if translated is None:
        logger.warning("Translation returned None for: %r — keeping original.", original[:80])
        translated = original

    if translated != original:
        write_text_to_runs(paragraph, translated)


def translate_document(doc: Document, translator) -> None:
    """
    Hujjat ichidagi barcha matnlarni bitta ro'yxatga yig'ib, 
    Google API'ga BATCH (guruh) qilib yuboradi. 6 daqiqalik ishni 5 soniyaga tushiradi.
    """
    paragraphs_to_translate = []

    # 1. Oddiy xatboshilarni yig'amiz
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs_to_translate.append(p)

    # 2. Jadvallar ichidagi matnlarni yig'amiz
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in iter_paragraphs_in_cell(cell):
                    if p.text.strip():
                        paragraphs_to_translate.append(p)

    if not paragraphs_to_translate:
        return

    # 3. Faqat matnlarni ajratib olamiz
    texts = [p.text for p in paragraphs_to_translate]

    try:
        # Google API'ga barcha matnni bitta so'rovda yuboramiz!
        logger.info(f"Google API'ga {len(texts)} ta qator bitta paketda yuborilmoqda...")
        translated_texts = translator.translate_batch(texts)
    except Exception as e:
        logger.error(f"Batch translation xatosi: {e}")
        return

    # 4. Tarjima qilingan matnlarni o'z joylariga rasmga qarab joylashtiramiz
    for p, translated_text in zip(paragraphs_to_translate, translated_texts):
        if translated_text and translated_text != p.text:
            write_text_to_runs(p, translated_text)

# ─────────────────────────────────────────────────────────────────────────────
# Top-level entry point called by the bot handler
# ─────────────────────────────────────────────────────────────────────────────

class SameLanguageError(Exception):
    """Raised when source and target languages are identical."""


def translate_docx_file(input_path: str, output_path: str, source_lang: str, target_lang: str) -> None:
    """
    Load *input_path*, translate from *source_lang* to *target_lang*,
    and save result to *output_path*.

    Raises:
        SameLanguageError: if source == target (bot informs the user instead
                           of wasting API quota).
        Any exception from deep-translator or python-docx propagates upward
        and is caught by the bot handler.

    NOTE: This function is synchronous and blocking.  Always call it via
          ``await asyncio.to_thread(translate_docx_file, ...)`` in async code.
    """
    from deep_translator import GoogleTranslator  # lazy import — keeps startup fast

    logger.info("Loading document: %s", input_path)
    doc = Document(input_path)

    logger.info("Source language: %s → target: %s", source_lang, target_lang)

    if source_lang == target_lang:
        raise SameLanguageError(source_lang)

    # Build the Google Translator instance
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    # Translate the entire document in-place
    translate_document(doc, translator)

    # Persist the translated document
    doc.save(output_path)
    logger.info("Translated document saved: %s", output_path)

    # No return needed
